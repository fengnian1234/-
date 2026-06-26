"""
周报生成服务 - 生成 DOCX 格式的口碑监控周报
使用 python-docx 生成紧凑优雅中式排版
v3.6: +趋势分析图表（matplotlib）
"""
import os
from collections import defaultdict
from datetime import datetime, timedelta
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# matplotlib 可选依赖，用于趋势图；不可用时降级为纯文字趋势
try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    import matplotlib.ticker as mticker
    _HAS_MPL = True
except ImportError:
    _HAS_MPL = False

from services.logger import info, warning, debug

from models import SessionLocal, PlatformMention
from services.monitor import (
    get_mentions_summary, search_platform_mentions,
    get_platform_review_links, MONITOR_PLATFORMS,
)
from config import BNB_NAME, BNB_ADDRESS, BNB_PHONE

# 输出目录
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REPORT_DIR = os.path.join(PROJECT_ROOT, "reports")

# 颜色常量
COLOR_PRIMARY = "2a3d33"      # 深绿
COLOR_SECONDARY = "5b7b6f"    # 中绿
COLOR_ACCENT = "8b7355"       # 暖棕
COLOR_MUTED = "6b6b66"        # 灰色
COLOR_POSITIVE = "388E3C"     # 正面绿（稍深）
COLOR_NEUTRAL = "f9a825"      # 琥珀黄
COLOR_NEGATIVE = "d32f2f"     # 深红
COLOR_BORDER = "d5cec0"       # 边框米色
COLOR_WHITE = "ffffff"
COLOR_DARK = "1a1a1a"
COLOR_CARD_BG = "faf8f5"      # 卡片暖白背景
COLOR_HEADER_BG = "e8e4db"    # 表头背景


def generate_weekly_report() -> dict:
    """主入口：收集最新数据并生成 DOCX 周报"""
    os.makedirs(REPORT_DIR, exist_ok=True)

    # 1. 收集最新平台数据
    info("📡 正在收集平台最新数据...")
    search_results = search_platform_mentions()

    # 2. 获取汇总
    summary = get_mentions_summary()

    # 3. 生成 DOCX
    now = datetime.utcnow()
    week_start = (now - timedelta(days=7)).strftime("%Y.%m.%d")
    week_end = now.strftime("%Y.%m.%d")
    filename = f"云上归墅_口碑周报_{now.strftime('%Y%m%d_%H%M')}.docx"
    filepath = os.path.join(REPORT_DIR, filename)

    doc = _build_docx(summary, search_results, week_start, week_end)
    doc.save(filepath)

    info(f"✅ 周报已生成: {filepath}")

    return {
        "success": True,
        "filename": filename,
        "filepath": filepath,
        "overall_rating": summary.get("overall_rating"),
        "total_mentions": sum(
            p.get("total", 0) for p in summary.get("platforms", {}).values()
        ),
        "platforms_covered": len(MONITOR_PLATFORMS),
        "week_range": f"{week_start} - {week_end}",
    }


# ═══════════════════════════ 辅助 ═══════════════════════════

def _set_cell_shading(cell, color: str):
    """设置表格单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_rich_run(paragraph, text: str, font_size: Pt = None, bold: bool = False,
                  color: str = None, italic: bool = False,
                  font_name: str = 'Microsoft YaHei'):
    """添加带样式的文本片段"""
    run = paragraph.add_run(text)
    if font_size:
        run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*_hex_to_rgb(color))
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return run


def _hex_to_rgb(hex_color: str) -> tuple:
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def _add_decorative_line(doc, color: str = COLOR_BORDER, width_pt: int = 1):
    """添加装饰性细线（用段落边框实现）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{width_pt * 8}" w:space="1" '
        f'w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _add_heading(doc: Document, text: str, level: int = 1):
    """添加带样式的标题（紧凑版）"""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(10) if level == 1 else Pt(6)
    heading.paragraph_format.space_after = Pt(4) if level == 1 else Pt(2)
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 1:
            run.font.color.rgb = RGBColor(*_hex_to_rgb(COLOR_PRIMARY))
            run.font.size = Pt(14)
        elif level == 2:
            run.font.color.rgb = RGBColor(*_hex_to_rgb(COLOR_SECONDARY))
            run.font.size = Pt(11)
    return heading


try:
    from PIL import Image as PILImage
    _HAS_PIL = True
except ImportError:
    _HAS_PIL = False


def _prepare_image_for_docx(full_path: str) -> tuple:
    """
    准备图片供 python-docx 嵌入。返回 (stream_or_path, width_inches)
    - 小红书下载的图片常为 WebP 格式但扩展名 .jpg，python-docx 不支持
    - 用 Pillow 检测格式，WebP 则转 JPEG 到 BytesIO
    - 同时用 Pillow 获取真实宽高比来确定图片宽度
    """
    img_width = Inches(2.1)  # 默认方形宽度

    if not os.path.exists(full_path):
        return (full_path, img_width)

    size_kb = os.path.getsize(full_path) / 1024

    if _HAS_PIL:
        try:
            img = PILImage.open(full_path)
            w, h = img.size
            aspect = w / h if h > 0 else 1.0
            if aspect < 0.75:
                img_width = Inches(1.6)
            elif aspect > 1.5:
                img_width = Inches(2.6)
            else:
                img_width = Inches(2.1)

            # 检测 WebP 格式并转换为 JPEG
            if img.format == 'WEBP':
                buf = BytesIO()
                img.convert('RGB').save(buf, format='JPEG', quality=90)
                buf.seek(0)
                img.close()
                return (buf, img_width)
            img.close()
        except Exception:
            # Pillow 失败时回退到文件大小启发式
            debug("Pillow 图片处理失败，回退到文件大小启发式", exc_info=True)
            if size_kb > 250:
                img_width = Inches(2.6)
            elif size_kb < 100:
                img_width = Inches(1.7)
    else:
        # 无 Pillow，用文件大小估算
        if size_kb > 250:
            img_width = Inches(2.6)
        elif size_kb < 100:
            img_width = Inches(1.7)

    return (full_path, img_width)


# ═══════════════════════════ 趋势分析 ═══════════════════════════

# matplotlib 中文字体配置
if _HAS_MPL:
    _CHINESE_FONTS = ['Microsoft YaHei', 'SimHei', 'WenQuanYi Micro Hei',
                       'Noto Sans CJK SC', 'PingFang SC', 'Heiti SC']
    _available_fonts = {f.name for f in fm.fontManager.ttflist}
    _selected_font = None
    for _fn in _CHINESE_FONTS:
        if _fn in _available_fonts:
            _selected_font = _fn
            break
    if not _selected_font:
        _selected_font = 'sans-serif'
    plt.rcParams['font.sans-serif'] = [_selected_font, 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False


def _query_weekly_trends(db, weeks: int = 6) -> list:
    """从 PlatformMention 查询最近 N 周的逐周趋势数据"""
    cutoff = datetime.utcnow() - timedelta(weeks=weeks)

    mentions = db.query(PlatformMention).filter(
        PlatformMention.collected_at >= cutoff
    ).order_by(PlatformMention.collected_at.asc()).all()

    # 按 ISO 周分组
    weekly = defaultdict(lambda: {"total": 0, "ratings": [], "positive": 0, "neutral": 0, "negative": 0})

    for m in mentions:
        if not m.collected_at:
            continue
        iso = m.collected_at.isocalendar()
        week_label = f"{iso[0]}W{iso[1]:02d}"
        weekly[week_label]["total"] += 1
        if m.rating and m.rating > 0:
            weekly[week_label]["ratings"].append(m.rating)
        sent = m.sentiment or "neutral"
        if sent == "positive":
            weekly[week_label]["positive"] += 1
        elif sent == "negative":
            weekly[week_label]["negative"] += 1
        else:
            weekly[week_label]["neutral"] += 1

    result = []
    for week_label in sorted(weekly.keys()):
        data = weekly[week_label]
        total = data["total"]
        ratings = data["ratings"]
        avg_rating = round(sum(ratings) / len(ratings), 2) if ratings else 0
        pos_pct = round(data["positive"] / total * 100, 1) if total > 0 else 0
        neu_pct = round(data["neutral"] / total * 100, 1) if total > 0 else 0
        neg_pct = round(data["negative"] / total * 100, 1) if total > 0 else 0

        result.append({
            "week": week_label,
            "total": total,
            "avg_rating": avg_rating,
            "positive": data["positive"],
            "neutral": data["neutral"],
            "negative": data["negative"],
            "pos_pct": pos_pct,
            "neu_pct": neu_pct,
            "neg_pct": neg_pct,
        })

    return result


def _generate_trend_chart(trend_data: list, chart_type: str, output_dir: str) -> str | None:
    """用 matplotlib 生成单张趋势图 PNG，返回文件路径"""
    if not _HAS_MPL or len(trend_data) < 2:
        return None

    weeks = [d["week"] for d in trend_data]
    # 配色与报告统一
    GREEN = '#2a3d33'
    MUTED = '#8b7355'
    POS = '#388E3C'
    NEU = '#f9a825'
    NEG = '#d32f2f'
    BG = '#faf8f5'

    fig, ax = plt.subplots(figsize=(6.8, 2.8))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)

    if chart_type == "rating":
        ratings = [d["avg_rating"] for d in trend_data]
        ax.plot(weeks, ratings, marker='o', color=GREEN, linewidth=2.2,
                markersize=7, markerfacecolor='white', markeredgewidth=2,
                markeredgecolor=GREEN)
        ax.set_ylabel('平均评分', fontsize=10, color=GREEN)
        ax.set_title('评分趋势', fontsize=12, fontweight='bold', color=GREEN, pad=8)
        ax.set_ylim(max(0, min(ratings) - 0.5), 5.1)
        for i, (x, y) in enumerate(zip(weeks, ratings)):
            ax.annotate(f'{y:.1f}', (x, y), textcoords="offset points",
                        xytext=(0, 10), ha='center', fontsize=8.5, color=GREEN)
    elif chart_type == "mentions":
        totals = [d["total"] for d in trend_data]
        bars = ax.bar(weeks, totals, color=GREEN, width=0.45, alpha=0.85, edgecolor='white', linewidth=0.5)
        ax.set_ylabel('提及数', fontsize=10, color=GREEN)
        ax.set_title('提及数量趋势', fontsize=12, fontweight='bold', color=GREEN, pad=8)
        for bar, val in zip(bars, totals):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                    str(val), ha='center', fontsize=8.5, color=GREEN)
    elif chart_type == "sentiment":
        x = range(len(weeks))
        width = 0.45
        pos_vals = [d["positive"] for d in trend_data]
        neu_vals = [d["neutral"] for d in trend_data]
        neg_vals = [d["negative"] for d in trend_data]
        ax.bar(x, pos_vals, width, label='正面', color=POS, alpha=0.85)
        ax.bar(x, neu_vals, width, bottom=pos_vals, label='中性', color=NEU, alpha=0.85)
        bottom_neu = [p + n for p, n in zip(pos_vals, neu_vals)]
        ax.bar(x, neg_vals, width, bottom=bottom_neu, label='负面', color=NEG, alpha=0.85)
        ax.set_xticks(x)
        ax.set_xticklabels(weeks)
        ax.set_ylabel('评价数', fontsize=10, color=GREEN)
        ax.set_title('情感分布趋势', fontsize=12, fontweight='bold', color=GREEN, pad=8)
        ax.legend(fontsize=8, loc='upper right', framealpha=0.9)

    # 通用设置
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#d5cec0')
    ax.spines['bottom'].set_color('#d5cec0')
    ax.tick_params(axis='both', colors='#6b6b66', labelsize=8.5)
    ax.yaxis.set_major_locator(mticker.MaxNLocator(integer=True, nbins=5))

    fig.tight_layout(pad=1.5)
    filepath = os.path.join(output_dir, f"trend_{chart_type}.png")
    fig.savefig(filepath, dpi=140, bbox_inches='tight', facecolor=BG, edgecolor='none')
    plt.close(fig)
    return filepath


def _prepare_trend_charts(db) -> tuple:
    """查询趋势数据并生成三张图表 PNG，返回 (chart_paths, trend_data)"""
    trend_data = _query_weekly_trends(db, weeks=6)
    if not _HAS_MPL or len(trend_data) < 2:
        return [], trend_data

    # 保存到临时目录
    chart_dir = os.path.join(REPORT_DIR, "_trend_charts")
    os.makedirs(chart_dir, exist_ok=True)

    charts = []
    for ctype in ("rating", "mentions", "sentiment"):
        fp = _generate_trend_chart(trend_data, ctype, chart_dir)
        if fp and os.path.exists(fp):
            charts.append(fp)

    return charts, trend_data


# ═══════════════════════════ 文档构建 ═══════════════════════════

def _build_docx(summary: dict, search_results: dict,
                week_start: str, week_end: str) -> Document:
    """构建 DOCX 文档 v3 — 紧凑优雅版"""
    doc = Document()

    # ── 页面设置（A4 标准边距）──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── 全局样式（紧凑）──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.2
    style.paragraph_format.space_after = Pt(3)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ─── 封面 ───
    _build_cover(doc, week_start, week_end)
    doc.add_page_break()

    # ─── 一、数据概览 ───
    _add_heading(doc, "一、本周数据概览", level=1)
    _add_decorative_line(doc)
    _build_dashboard(doc, summary, week_start, week_end)

    # ─── 二、平台评价 & 图片明细（合并原二+三）───
    _add_heading(doc, "二、平台评价 & 图片明细", level=1)
    _add_decorative_line(doc)
    _build_platform_with_images(doc, summary)

    # ─── 三、舆情分析 ───
    _add_heading(doc, "三、舆情分析与改进建议", level=1)
    _add_decorative_line(doc)
    _build_analysis(doc, summary)

    # ─── 四、趋势分析（图表）───
    db2 = SessionLocal()
    try:
        trend_charts, trend_data = _prepare_trend_charts(db2)
    finally:
        db2.close()

    if trend_charts or trend_data:
        _add_heading(doc, "四、趋势分析", level=1)
        _add_decorative_line(doc)

        if trend_charts:
            intro = doc.add_paragraph()
            intro.paragraph_format.space_after = Pt(6)
            _add_rich_run(intro,
                f"以下图表展示最近 {len(trend_data)} 周的口碑变化趋势（数据截止 {datetime.utcnow().strftime('%m.%d')}）：",
                Pt(10.5), color=COLOR_DARK)

            # 嵌入三张趋势图（每张图配简短解读）
            chart_labels = {
                "trend_rating.png": "📈 评分趋势 — 反映客人满意度变化，持续上升表明服务改进有效",
                "trend_mentions.png": "📊 提及数量 — 反映品牌曝光度波动，受旅游淡旺季和内容运营影响",
                "trend_sentiment.png": "🎭 情感分布 — 正面(绿)/中性(黄)/负面(红)占比变化，直观显示口碑健康度",
            }
            for chart_path in trend_charts:
                fname = os.path.basename(chart_path)
                label_text = chart_labels.get(fname, "")

                # 趋势图宽度 5.8"，居中
                chart_p = doc.add_paragraph()
                chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_p.paragraph_format.space_before = Pt(8)
                chart_p.paragraph_format.space_after = Pt(2)

                try:
                    prepared_path, _ = _prepare_image_for_docx(chart_path)
                    run = chart_p.add_run()
                    run.add_picture(prepared_path, width=Inches(5.8))
                except Exception as e:
                    _add_rich_run(chart_p, f"[图表加载失败]", Pt(8), color=COLOR_NEGATIVE)
                    warning(f"嵌入趋势图失败 {chart_path}: {e}")
                    continue

                if label_text:
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(6)
                    _add_rich_run(caption, label_text, Pt(8.5), color=COLOR_MUTED)
        else:
            # 无图表时的纯文字趋势
            trend_note = doc.add_paragraph()
            trend_note.paragraph_format.space_after = Pt(6)
            if len(trend_data) < 2:
                _add_rich_run(trend_note,
                    "⚠ 趋势数据不足（需至少2周历史数据），将在后续周报中自动展示评分/提及/情感趋势图。",
                    Pt(10.5), color=COLOR_MUTED)
            elif not _HAS_MPL:
                _add_rich_run(trend_note,
                    "⚠ matplotlib 未安装，无法生成趋势图。安装后即可自动展示趋势图。",
                    Pt(10.5), color=COLOR_MUTED)

            # 文字趋势概述
            if len(trend_data) >= 2:
                latest = trend_data[-1]
                prev = trend_data[-2]
                dr = latest.get("avg_rating", 0) - prev.get("avg_rating", 0)
                dt = latest.get("total", 0) - prev.get("total", 0)

                text_trend = doc.add_paragraph()
                text_trend.paragraph_format.left_indent = Cm(0.3)
                text_trend.paragraph_format.space_after = Pt(2)
                direction_r = "↑" if dr > 0 else ("↓" if dr < 0 else "→")
                direction_t = "↑" if dt > 0 else ("↓" if dt < 0 else "→")
                color_r = COLOR_POSITIVE if dr >= 0 else COLOR_NEGATIVE
                color_t = COLOR_POSITIVE if dt >= 0 else COLOR_NEGATIVE
                _add_rich_run(text_trend, f"与上周相比：评分 {latest.get('avg_rating',0)} ", Pt(10), color=COLOR_DARK)
                _add_rich_run(text_trend, f"{direction_r}{abs(dr):.1f}", Pt(10), bold=True, color=color_r)
                _add_rich_run(text_trend, f"  ·  提及数 {latest.get('total',0)} ", Pt(10), color=COLOR_DARK)
                _add_rich_run(text_trend, f"{direction_t}{abs(dt)}", Pt(10), bold=True, color=color_t)

    # ─── 页脚 ───
    _add_decorative_line(doc, COLOR_BORDER, 1)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(4)
    _add_rich_run(footer, f"—— {BNB_NAME} · 自动生成 ——  {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", Pt(8), color=COLOR_MUTED)

    return doc


# ═══════════════════════════ 封面 ═══════════════════════════

def _build_cover(doc, week_start: str, week_end: str):
    """紧凑封面 — 约1/3页"""
    for _ in range(3):
        doc.add_paragraph()

    _add_decorative_line(doc, COLOR_ACCENT, 2)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_rich_run(title, BNB_NAME, Pt(26), bold=True, color=COLOR_PRIMARY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(4)
    _add_rich_run(subtitle, "口碑监控周报", Pt(17), color=COLOR_SECONDARY)

    # 期数
    week_num = datetime.utcnow().isocalendar()[1]
    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period.paragraph_format.space_before = Pt(6)
    _add_rich_run(period, f"2026年 · 第{week_num}周", Pt(10), color=COLOR_MUTED)

    _add_decorative_line(doc, COLOR_ACCENT, 2)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(10)

    # 报告周期
    date_info = doc.add_paragraph()
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_rich_run(date_info, f"{week_start}  —  {week_end}", Pt(11), color=COLOR_MUTED)

    # 地址电话同行
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(4)
    _add_rich_run(contact, f"{BNB_ADDRESS}    ☎ {BNB_PHONE}", Pt(9), color=COLOR_MUTED)


# ═══════════════════════════ 仪表盘 ═══════════════════════════

def _build_dashboard(doc, summary: dict, week_start: str, week_end: str):
    """紧凑仪表盘 — KPI卡片 + 评分表"""
    overall = summary.get("overall_rating")
    platforms_data = summary.get("platforms", {})

    total_mentions = sum(p.get("total", 0) for p in platforms_data.values())
    total_pos = sum(p.get("positive", 0) for p in platforms_data.values())
    total_neu = sum(p.get("neutral", 0) for p in platforms_data.values())
    total_neg = sum(p.get("negative", 0) for p in platforms_data.values())
    total_all = total_pos + total_neu + total_neg
    active_platforms = sum(1 for p in platforms_data.values() if p.get("total", 0) > 0)

    # 摘要（紧凑）
    summary_text = (
        f"监控 {len(MONITOR_PLATFORMS)} 个平台，{active_platforms} 个有活跃数据，"
        f"累计 {total_mentions} 条评价。"
    )
    if overall:
        summary_text += f"综合评分 ⭐{overall}/5.0。"
    p = doc.add_paragraph(summary_text)
    p.paragraph_format.space_after = Pt(6)

    # KPI 卡片（单行4列，每列独立底色）
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    card_colors = ["f0ede6", "e8f0ec", "eaf0f0", "f0f0e8"]
    kpis = [
        ("综合评分", f"⭐{overall}/5.0" if overall else "暂无", COLOR_ACCENT),
        ("活跃平台", str(active_platforms), COLOR_PRIMARY),
        ("评价总数", str(total_mentions), COLOR_SECONDARY),
        ("好评率", f"{round(total_pos/total_all*100)}%" if total_all > 0 else "暂无",
         COLOR_POSITIVE if (total_all > 0 and total_pos/total_all >= 0.7) else COLOR_NEUTRAL),
    ]

    for i, (label, value, color) in enumerate(kpis):
        cell = kpi_table.rows[0].cells[i]
        cell.width = Cm(3.6)
        _set_cell_shading(cell, card_colors[i])
        cell.paragraphs[0].clear()

        p_val = cell.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_val.paragraph_format.space_before = Pt(4)
        p_val.paragraph_format.space_after = Pt(0)
        _add_rich_run(p_val, value, Pt(16), bold=True, color=color)

        p_label = cell.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_label.paragraph_format.space_before = Pt(1)
        p_label.paragraph_format.space_after = Pt(4)
        _add_rich_run(p_label, label, Pt(8.5), color=COLOR_MUTED)

    # 评分概览表
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    _add_rich_run(p, "各平台评分概览", Pt(10.5), bold=True, color=COLOR_PRIMARY)

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["平台", "条数", "均分", "评分", "👍", "😐", "👎"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _add_rich_run(p, header, Pt(8.5), bold=True, color=COLOR_DARK)
        _set_cell_shading(cell, COLOR_HEADER_BG)

    for platform in MONITOR_PLATFORMS:
        data = platforms_data.get(platform, {})
        total = data.get("total", 0)
        avg = data.get("avg_rating")
        row = table.add_row()

        vals = [
            platform,
            str(total),
            f"{avg}/5.0" if avg else "—",
            ("★" * round(avg) + "☆" * (5 - round(avg))) if avg else "—",
            str(data.get("positive", 0)),
            str(data.get("neutral", 0)),
            str(data.get("negative", 0)),
        ]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            color = COLOR_DARK
            if ci == 4:
                color = COLOR_POSITIVE
            elif ci == 5:
                color = COLOR_NEUTRAL
            elif ci == 6:
                color = COLOR_NEGATIVE
            _add_rich_run(p, val, Pt(8.5), bold=(ci in (4, 5, 6)), color=color)


# ═══════════════════════════ 平台评价 & 图片（合并版）═══════════

def _build_platform_with_images(doc, summary: dict):
    """合并平台详情 + 评价图片为单一紧凑区域"""
    platforms_data = summary.get("platforms", {})

    # 按均分降序
    sorted_platforms = sorted(
        [(p, d) for p, d in platforms_data.items()],
        key=lambda x: x[1].get("avg_rating") or 0,
        reverse=True,
    )

    # 从 DB 获取所有有图片的 mention
    db = SessionLocal()
    try:
        all_img_mentions = db.query(PlatformMention).filter(
            PlatformMention.local_images != None
        ).order_by(PlatformMention.collected_at.desc()).all()

        # 按平台分组
        imgs_by_platform = defaultdict(list)
        for m in all_img_mentions:
            if m.local_images:
                imgs_by_platform[m.platform].append(m)
    finally:
        db.close()

    for platform, data in sorted_platforms:
        total = data.get("total", 0)
        avg = data.get("avg_rating")
        pos, neu, neg = data.get("positive", 0), data.get("neutral", 0), data.get("negative", 0)

        # ── 平台标题行（紧凑：标题+评分+指标同行）──
        _add_heading(doc, f"▸ {platform}", level=2)

        if total == 0:
            p = doc.add_paragraph(f"  本周暂无数据。")
            p.paragraph_format.left_indent = Cm(0.3)
            _add_rich_run(p, "", Pt(8), color=COLOR_MUTED)
            _add_decorative_line(doc, COLOR_BORDER, 1)
            continue

        # 指标行（单行显示所有关键数据）
        stars = "★" * round(avg) + "☆" * (5 - round(avg)) if avg else "—"
        rating_text = f"{avg}/5.0" if avg else "—"
        indicators = doc.add_paragraph()
        indicators.paragraph_format.left_indent = Cm(0.2)
        indicators.paragraph_format.space_after = Pt(2)
        _add_rich_run(indicators, f"📊 {total}条评价  |  ", Pt(9.5), color=COLOR_MUTED)
        _add_rich_run(indicators, f"均分 {rating_text} {stars}  |  ", Pt(9.5), color=COLOR_DARK)
        _add_rich_run(indicators, f"👍{pos}  ", Pt(9.5), color=COLOR_POSITIVE)
        _add_rich_run(indicators, f"😐{neu}  ", Pt(9.5), color=COLOR_NEUTRAL)
        _add_rich_run(indicators, f"👎{neg}", Pt(9.5), color=COLOR_NEGATIVE)

        # 情感条形
        all_count = pos + neu + neg
        if all_count > 0:
            bar_line = doc.add_paragraph()
            bar_line.paragraph_format.left_indent = Cm(0.3)
            bar_line.paragraph_format.space_before = Pt(1)
            bar_line.paragraph_format.space_after = Pt(2)
            pos_w = max(round(pos / all_count * 24), 1 if pos > 0 else 0)
            neg_w = max(round(neg / all_count * 24), 1 if neg > 0 else 0)
            neu_w = max(24 - pos_w - neg_w, 0)
            _add_rich_run(bar_line, "█" * pos_w, Pt(7), color=COLOR_POSITIVE)
            _add_rich_run(bar_line, "█" * neu_w, Pt(7), color=COLOR_NEUTRAL)
            _add_rich_run(bar_line, "█" * neg_w, Pt(7), color=COLOR_NEGATIVE)
            _add_rich_run(bar_line, f"  {pos}好评 / {neu}中评 / {neg}差评", Pt(7.5), color=COLOR_MUTED)

        # ── 最新评价摘要（仅文字）──
        latest = data.get("latest", {})
        if latest and isinstance(latest, dict):
            title = latest.get("title", "")
            content = latest.get("content", "")
            rating_val = latest.get("rating")
            url = latest.get("url", "")
            sentiment = latest.get("sentiment", "neutral")
            sent_emoji = {"positive": "😊", "neutral": "😐", "negative": "😞"}.get(sentiment, "")

            if title or content:
                quote_box = doc.add_paragraph()
                quote_box.paragraph_format.left_indent = Cm(0.5)
                quote_box.paragraph_format.space_before = Pt(3)
                quote_box.paragraph_format.space_after = Pt(2)
                if rating_val:
                    _add_rich_run(quote_box, f"【{rating_val}分】", Pt(9), bold=True, color=COLOR_ACCENT)
                snippet = (content or title)[:180]
                _add_rich_run(quote_box, f' {sent_emoji} "{snippet}..."', Pt(9), italic=True, color=COLOR_DARK)
            if url:
                link_p = doc.add_paragraph()
                link_p.paragraph_format.left_indent = Cm(0.5)
                link_p.paragraph_format.space_after = Pt(4)
                _add_rich_run(link_p, f"🔗 {url[:130]}", Pt(7.5), color=COLOR_SECONDARY)

        # ── 评价配图（图片网格：2列表格）──
        platform_mentions = imgs_by_platform.get(platform, [])
        if platform_mentions:
            # 收集该平台所有可用图片路径
            all_local_imgs = []
            for mention in platform_mentions[:6]:  # 最多取6条有图的评价
                for img_path in (mention.local_images or []):
                    full_path = os.path.join(PROJECT_ROOT, img_path)
                    if os.path.exists(full_path) and os.path.getsize(full_path) > 500:
                        all_local_imgs.append((full_path, mention.title or "", mention.rating))
                        if len(all_local_imgs) >= 8:  # 每平台最多8张图（4行×2列）
                            break
                if len(all_local_imgs) >= 8:
                    break

            if all_local_imgs:
                img_label = doc.add_paragraph()
                img_label.paragraph_format.left_indent = Cm(0.2)
                img_label.paragraph_format.space_before = Pt(4)
                img_label.paragraph_format.space_after = Pt(2)
                _add_rich_run(img_label, f"📷 评价配图（{len(all_local_imgs)}张）：", Pt(9), color=COLOR_MUTED)

                # 图片网格表格（每行2列）
                cols = 2
                rows = (len(all_local_imgs) + cols - 1) // cols
                img_table = doc.add_table(rows=rows, cols=cols)
                img_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for idx, (img_path, img_title, img_rating) in enumerate(all_local_imgs):
                    row_idx = idx // cols
                    col_idx = idx % cols
                    cell = img_table.rows[row_idx].cells[col_idx]

                    # 清除默认段落
                    cell.paragraphs[0].clear()
                    cell_p = cell.paragraphs[0]
                    cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell_p.paragraph_format.space_before = Pt(3)
                    cell_p.paragraph_format.space_after = Pt(1)

                    # 准备图片（WebP→JPEG 转换 + 智能宽度）
                    img_src, img_width = _prepare_image_for_docx(img_path)

                    try:
                        run = cell_p.add_run()
                        run.add_picture(img_src, width=img_width)
                    except Exception as e:
                        _add_rich_run(cell_p, f"[图片加载失败]", Pt(7), color=COLOR_NEGATIVE)
                        warning(f"嵌入图片失败 {img_path}: {e}")
                        continue

                    # 图片下方小字标注
                    caption_p = cell.add_paragraph()
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_p.paragraph_format.space_before = Pt(1)
                    caption_p.paragraph_format.space_after = Pt(2)
                    stars_str = f" {'★' * round(img_rating)}{'☆' * (5 - round(img_rating))}" if img_rating else ""
                    _add_rich_run(caption_p, f"{img_title[:20]}{stars_str}", Pt(7), color=COLOR_MUTED)

                # 移除未填充的行
                # (python-docx 不支持直接删行，空单元格显示为空白即可)

        # 平台间分隔
        _add_decorative_line(doc, COLOR_BORDER, 1)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

    # ── 各平台评价入口 ──
    _add_heading(doc, "各平台评价入口", level=2)
    platforms_links = get_platform_review_links()
    for platform_name, info_item in platforms_links.items():
        link_p = doc.add_paragraph()
        link_p.paragraph_format.left_indent = Cm(0.3)
        link_p.paragraph_format.space_after = Pt(1)
        _add_rich_run(link_p, f"{info_item['icon']} {info_item['name']}：", Pt(9), bold=True, color=COLOR_MUTED)
        _add_rich_run(link_p, info_item['review_url'], Pt(8), color=COLOR_SECONDARY)


# ═══════════════════════════ 舆情分析 ═══════════════════════════

def _build_analysis(doc, summary: dict):
    """紧凑舆情分析与建议"""
    platforms_data = summary.get("platforms", {})

    total_pos = sum(p.get("positive", 0) for p in platforms_data.values())
    total_neu = sum(p.get("neutral", 0) for p in platforms_data.values())
    total_neg = sum(p.get("negative", 0) for p in platforms_data.values())
    total_all = total_pos + total_neu + total_neg

    if total_all > 0:
        pos_pct = round(total_pos / total_all * 100)

        doc.add_paragraph()
        _add_heading(doc, "情感分布总览", level=2)

        analysis = doc.add_paragraph()
        analysis.paragraph_format.space_after = Pt(4)
        _add_rich_run(analysis,
            f"本周共 {total_all} 条有效评价——正面 {total_pos} 条（{pos_pct}%），"
            f"中性 {total_neu} 条，负面 {total_neg} 条。",
            Pt(10.5), color=COLOR_DARK)

        # 健康度评估
        assessment = doc.add_paragraph()
        assessment.paragraph_format.left_indent = Cm(0.3)
        assessment.paragraph_format.space_after = Pt(4)
        if pos_pct >= 80:
            _add_rich_run(assessment, "🟢 口碑健康度：优秀  ", Pt(10.5), bold=True, color=COLOR_POSITIVE)
            _add_rich_run(assessment,
                "整体口碑优秀，客人满意度高。保持现有服务标准，持续关注新增评价，差评48小时内回复。", Pt(10), color=COLOR_DARK)
        elif pos_pct >= 60:
            _add_rich_run(assessment, "🟡 口碑健康度：良好  ", Pt(10.5), bold=True, color=COLOR_NEUTRAL)
            _add_rich_run(assessment,
                "口碑良好但仍有改进空间。梳理差评共性原因，制定改进计划，积极回复各平台评价。", Pt(10), color=COLOR_DARK)
        else:
            _add_rich_run(assessment, "🔴 口碑健康度：需关注  ", Pt(10.5), bold=True, color=COLOR_NEGATIVE)
            _add_rich_run(assessment,
                "口碑需重点改善。立即排查差评根因，召开服务质量会议，制定系统化改进方案。", Pt(10), color=COLOR_DARK)

        if total_neg > 0:
            doc.add_paragraph()
            _add_heading(doc, "⚠ 重点关注", level=2)
            neg_note = doc.add_paragraph()
            neg_note.paragraph_format.left_indent = Cm(0.3)
            _add_rich_run(neg_note,
                f"本周发现 {total_neg} 条负面评价，建议逐条分析并48小时内回复处理。回复「收集口碑」获取最新数据。",
                Pt(10), color=COLOR_NEGATIVE)
    else:
        p = doc.add_paragraph("本周暂无足够数据进行舆情分析。")
        p.paragraph_format.first_line_indent = Cm(0.74)

    # 改进建议（紧凑列表）
    doc.add_paragraph()
    _add_heading(doc, "下周改进建议", level=2)

    suggestions = [
        ("评价回复", "每日检查携程/美团/飞猪最新评价，差评48小时内真诚回复"),
        ("好评引导", "退房时引导满意客人在携程/美团写评价（退房好评推送已自动执行）"),
        ("内容运营", "小红书和抖音每周发布2-3条民宿日常/周边攻略内容，提升自然曝光量"),
        ("竞品对标", "关注庐山区域同类精品民宿评价趋势，了解差异化服务方向"),
        ("服务优化", "针对高频差评关键词（隔音/卫生/水温）做预防性检查和改进"),
    ]

    for i, (title, detail) in enumerate(suggestions, 1):
        item_p = doc.add_paragraph()
        item_p.paragraph_format.left_indent = Cm(0.2)
        item_p.paragraph_format.space_before = Pt(1)
        item_p.paragraph_format.space_after = Pt(1)
        _add_rich_run(item_p, f"{i}. {title}：", Pt(9.5), bold=True, color=COLOR_PRIMARY)
        _add_rich_run(item_p, detail, Pt(9.5), color=COLOR_DARK)
