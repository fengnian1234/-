"""
云上归墅专属项目规划方案 DOCX -- 面向民宿方展示项目功能与前景
系统未落成视角：展示的是"我们将为您建设什么"而非"已经建成了什么"
纯文本安全字符，无高位 Unicode
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

DOC = Document()
for sec in DOC.sections:
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

PRI=RGBColor(0x4A,0x7C,0x59); PRID=RGBColor(0x3B,0x5E,0x47); ACC=RGBColor(0xC8,0x96,0x6A)
W=RGBColor(0xFF,0xFF,0xFF); DK=RGBColor(0x2A,0x33,0x28)
GR=RGBColor(0x5A,0x65,0x58); LG=RGBColor(0x8A,0x95,0x86); BD=RGBColor(0xD8,0xE0,0xD4)

def dline(typ="thick"):
    p=DOC.add_paragraph(); pf=p.paragraph_format
    pf.space_after=Pt(0); pf.space_before=Pt(0)
    c="B8D0B8" if typ=="thick" else "C8966A"; s="12" if typ=="thick" else "4"
    pPr=p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:bottom w:val="single" w:sz="{s}" w:space="1" w:color="{c}"/></w:pBdr>'))

def stitle(tx):
    p=DOC.add_paragraph(); pf=p.paragraph_format; pf.space_before=Pt(18); pf.space_after=Pt(0)
    dline("thick")
    p2=DOC.add_paragraph(); pf2=p2.paragraph_format; pf2.space_after=Pt(4); pf2.space_before=Pt(4)
    r=p2.add_run(tx); r.font.size=Pt(20); r.font.color.rgb=PRID; r.bold=True
    dline("thin"); DOC.add_paragraph()

def body(tx,indent=False,sz=11,c=DK):
    p=DOC.add_paragraph(); pf=p.paragraph_format; pf.space_after=Pt(5)
    if indent: pf.left_indent=Cm(0.8)
    r=p.add_run(tx); r.font.size=Pt(sz); r.font.color.rgb=c

def bullet(ic,tx,sz=10.5):
    p=DOC.add_paragraph(); pf=p.paragraph_format; pf.space_after=Pt(3); pf.left_indent=Cm(0.6)
    r=p.add_run(f"{ic} {tx}"); r.font.size=Pt(sz); r.font.color.rgb=DK

def subsec(tx,sz=14):
    p=DOC.add_paragraph(); pf=p.paragraph_format; pf.space_after=Pt(4); pf.space_before=Pt(10)
    r=p.add_run(tx); r.font.size=Pt(sz); r.font.color.rgb=PRI; r.bold=True

# ================================================================
# 封面
# ================================================================
for _ in range(5):
    p=DOC.add_paragraph(); pf=p.paragraph_format; pf.space_after=Pt(0)

p=DOC.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("云上 . 归墅"); r.font.size=Pt(32); r.font.color.rgb=PRID; r.bold=True
p2=DOC.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
pf2=p2.paragraph_format; pf2.space_after=Pt(4)
r2=p2.add_run("微信公众号智能服务系统 . 项目规划方案"); r2.font.size=Pt(16); r2.font.color.rgb=ACC
p3=DOC.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
pf3=p3.paragraph_format; pf3.space_after=Pt(8)
r3=p3.add_run("-"*36); r3.font.size=Pt(10); r3.font.color.rgb=BD
p4=DOC.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
r4=p4.add_run("云无心以出岫，鸟倦飞而知还"); r4.font.size=Pt(14); r4.font.color.rgb=ACC
DOC.add_paragraph()
p5=DOC.add_paragraph(); p5.alignment=WD_ALIGN_PARAGRAPH.CENTER
r5=p5.add_run("庐山山上 . 大林沟路27号  |  11间客房  |  2016年开业"); r5.font.size=Pt(11); r5.font.color.rgb=GR
p6=DOC.add_paragraph(); p6.alignment=WD_ALIGN_PARAGRAPH.CENTER
r6=p6.add_run("携程 4.7 分 . 85 条住客评价  |  2026年6月"); r6.font.size=Pt(10); r6.font.color.rgb=LG
DOC.add_page_break()

# ================================================================
# 1. 项目概述
# ================================================================
stitle("项目概述")

body("本方案为云上归墅民宿量身定制一套微信公众号智能服务系统。系统建成后，将完整覆盖住客「预订前 -- 入住中 -- 离店后」全周期的线上服务需求。住客关注公众号即可通过消息对话或菜单页面，自助获取房型展示、在线点单、游玩攻略、快捷服务和积分会员等全部功能。")

body("整套系统以 AI 智能管家为核心驱动，搭配 10 个移动端 H5 功能页面、积分会员体系和运营数据看板，为民宿的线上化服务能力带来质的提升。")

subsec("云上归墅 -- 庐山山上的精品民宿")
body("云上归墅位于庐山山上大林沟路 27 号，是一座 U 型三层山居小院，拥有 11 间精品客房（8 种房型），一楼配套「三山二水」咖啡馆。2016 年开业至今，在携程等 OTA 平台累计获得 4.7 分高评分和 85 条真实住客评价。", indent=True)

subsec("项目规划范围")
for it in ["11 间客房 . 8 种精品房型在线展示",
           "AI 智能管家 -- 三阶段全天候自动服务",
           "10 个移动端 H5 功能页面覆盖全部服务场景",
           "21 项快捷服务一键呼叫",
           "4 大 OTA 预订平台一键跳转对接",
           "积分 + 会员等级体系激励回头和好评",
           "运营数据看板 . 多平台订单统一管理",
           "自动周报生成 + 多渠道口碑监控"]:
    bullet("-", it)

DOC.add_page_break()

# ================================================================
# 2. AI 智能管家 -- 系统核心亮点
# ================================================================
stitle("AI 智能管家 -- 系统核心亮点")

body("AI 管家是本套系统最具差异化的核心模块。它能够感知客人所处的阶段，自动切换服务身份：预订前是热情专业的旅行顾问，预订后是贴心周到的专属管家，离店后是温暖持续的复购关怀。")

body("对民宿运营而言，一位 AI 管家可以同时响应所有住客的消息，全天候在线，不会疲惫、不会遗漏、不会情绪化，确保每一位客人都获得一致的优质服务体验。民宿前台无需夜间值守，重复性咨询由 AI 自动处理，人力成本显著降低。")

subsec("三阶段智能服务设计")

phases=[("阶段一：预订前 -- 免费旅行顾问",
         "任何关注公众号的用户均可使用。AI 以庐山本地旅行顾问身份，热情解答景点、路线、交通、美食等问题。在提供有用信息的过程中，自然展示民宿特色和房型亮点，引导用户产生预订兴趣。",
         ["零门槛使用，降低咨询心理负担",
          "在回答中自然融入民宿特色介绍",
          "引导用户跳转 OTA 平台查看和预订",
          "即使未入住也留下专业贴心的品牌印象"]),
        ("阶段二：预订后 -- 专属 AI 管家（核心体验）",
         "住客通过携程/美团/飞猪/大众点评预订后，在公众号内确认绑定即可解锁全部 AI 管家功能。AI 管家认识每位住客的姓名和入住信息，提供个性化一对一服务。这是整套系统体验最突出的阶段。",
         ["全天候即时响应，前台夜间无需值守",
          "21 项快捷服务一键呼叫直达员工",
          "实时推荐周边美食和游玩路线",
          "个性化关怀：生日月积分翻倍等贴心设计"]),
        ("阶段三：离店后 -- 持续复购关怀",
         "客人退房后，AI 自动切换为关怀模式，持续维护品牌连接，提升复购率。",
         ["退房后自动推送好评提醒（措辞温暖、不压迫）",
          "积分累计和会员等级升级自动通知",
          "会员等级越高回住折扣越大",
          "无需专人跟踪维护，系统自动运行"])]

for ti,de,its in phases:
    subsec(f">> {ti}")
    body(de,indent=True)
    for it in its: bullet("-", it)
    body("")

DOC.add_page_break()

# ================================================================
# 3. 页面功能规划
# ================================================================
stitle("移动端页面功能规划")

body("以下 10 个 H5 页面将嵌入微信公众号菜单，适配各类手机屏幕，支持深色/浅色模式自动切换，并可随春夏秋冬三季变换主题配色。所有页面均可直接在微信内打开，住客无需跳出微信。")

subsec("房型展示")
body("8 种精品房型在线展示。每间房将包含多张实拍图片、详细面积/床型/设施信息，价格透明标注。图片支持左右滑动轮播。住客可自助了解所有房型，大幅减少「有什么房型」「多少钱」「能住几人」等重复咨询。",indent=True)
bullet("-", "减少前台重复咨询量，释放人力")
bullet("-", "引导住客跳转 OTA 平台直接预订")

subsec("三山二水咖啡 -- 在线点单")
body("住客在房间内用手机浏览菜单、定制口味偏好、微信支付下单，送餐到房。饮品支持糖度选择、冰量选择，云雾茶支持红/绿茶和杯/壶两档规格。",indent=True)
bullet("-", "19 款产品覆盖咖啡、茶饮、简餐、甜品、特调")
bullet("-", "拓展民宿非房费收入来源")
bullet("-", "微信支付直接闭环，无需额外支付开发")

subsec("游玩攻略 + 美食探店")
body("将收录 5 条真实验证的庐山深度路线和 8 家周边美食推荐。数据来源于小红书、大众点评、携程的真实探店笔记，含详细图文介绍、必点菜和地图导航链接。",indent=True)
bullet("-", "大幅减少「附近有什么好玩的/好吃的」类咨询")
bullet("-", "可直接跳转高德/腾讯地图一键导航")

subsec("21 项快捷服务")
body("住客通过微信回复关键词或点击页面按钮即可呼叫服务，系统自动通知员工处理。覆盖客房服务、前台服务和设施维修三大类。",indent=True)
bullet("-", "客房：续住 . 打扫 . 补充用品 . 送餐 . 洗衣")
bullet("-", "前台：行李寄存 . 旅游咨询 . 叫车 . 叫醒 . 登山杖租借 . 特产代购等")
bullet("-", "维修：设施报修 . 空调 . 热水")

subsec("积分会员体系")
body("入住消费、每日签到、写平台评价、邀请好友均可获取积分。积分可兑换咖啡、房型升级、延迟退房和房费抵扣券。会员分银卡/金卡/钻石卡三级，等级越高折扣越大。生日当月积分自动 x1.5 倍。",indent=True)
bullet("-", "激励好评：写 OTA 评价 +80 分，助力提升平台评分")
bullet("-", "激励回头：积分累积 + 等级进阶推动复购")
bullet("-", "生日月 x1.5 倍积分 -- 温暖而不刻意的惊喜设计")

DOC.add_page_break()

# ================================================================
# 4. 运营管理后台
# ================================================================
stitle("运营管理后台")

body("系统将为云上归墅提供一套简洁实用的运营管理后台，老板在手机上即可查看全部数据，无需电脑，不受地点限制。")

subsec("多平台订单聚合看板")
body("携程/美团/飞猪/大众点评/小红书/抖音/直接预订 -- 7 个渠道的订单将统一展示管理。今日入住、今日退房、当前在住一目了然，未来 7 天到店预测辅助排班和物资准备。",indent=True)

subsec("服务请求追踪看板")
body("住客通过公众号发起的每一条服务请求自动记录，按待处理/处理中/已完成状态流转。告别纸条和对讲机，不再遗漏任何客人需求。",indent=True)

subsec("自动化运营工具")
body("系统将实现退房后定时推送好评提醒、每周自动生成经营周报、多渠道口碑持续监控等自动化运营功能。日常运营事务由系统自动完成，民宿团队只需查看结果、专注提升服务质量。",indent=True)

DOC.add_page_break()

# ================================================================
# 5. 预订渠道对接
# ================================================================
stitle("预订渠道对接")

body("云上归墅在以下主流 OTA 平台均可在线预订。系统将为每个平台配置专属跳转链接，住客在公众号内即可一键跳转至对应平台的预订页面。")
body("")

plats=[("携程旅行","搜索「云上归墅」","携程 4.7 分 . 85 条真实评价"),
       ("美团民宿","搜索「云上归墅」","民宿频道优选推荐"),
       ("飞猪旅行","搜索「云上归墅」","支付宝小程序直接预订"),
       ("大众点评","搜索「云上归墅」","真实住客评价参考")]

for nm,sc,no in plats:
    subsec(f">> {nm}")
    body(f"{sc}  --  {no}",indent=True)

body("")
body("预订成功后，住客在公众号内回复「绑定预订」，经前台确认即可解锁 AI 管家全部专属功能。",indent=True)
body("预订咨询电话：16607927666  .  前台服务时间：每日 8:00 - 22:00")

DOC.add_page_break()

# ================================================================
# 6. 住客使用方式
# ================================================================
stitle("住客使用方式（系统建成后）")

body("住客无需下载任何 App。在微信中搜索「云上归墅民宿」关注公众号，即可使用全部功能。以下是两种主要交互方式：")

subsec("方式一：回复关键词")
body("在公众号对话框发送数字或中文关键词，系统自动匹配并返回相应服务。所有回复均在 1 秒内响应。",indent=True)
for it in ["回复 1 或「房型」-- 查看所有房型图文介绍",
           "回复 2 或「菜单」-- 浏览咖啡简餐菜单并下单",
           "回复 3 或「攻略」-- 获取游玩路线推荐",
           "回复 4 或「服务」-- 查看 21 项快捷服务列表",
           "回复「美食」-- 周边美食深度探店推荐",
           "回复「地图」或「位置」-- 民宿定位与一键导航",
           "预订后回复任意自然语言问题 -- AI 管家智能应答"]:
    bullet("-", it)

subsec("方式二：公众号底部菜单")
body("点击公众号底部菜单栏，可直接进入 H5 页面浏览房型、在线点单、游玩攻略、快捷服务和积分中心。所有页面针对手机屏幕优化，支持深色/浅色模式自动切换，并随季节变换主题配色。",indent=True)

DOC.add_page_break()

# ================================================================
# 7. 系统页面清单
# ================================================================
stitle("系统页面清单（全部 11 个页面）")

pages=[("[首页] 民宿概览","精选房型推荐 . AI 管家入口 . 季节主题切换"),
       ("[房型] 房型列表","8 种房型卡片式展示 . 价格/人数/面积一览"),
       ("[房型] 房型详情","多图滑动轮播 . 设施清单 . 预订电话一键拨打"),
       ("[点单] 在线点单","4 大品类菜单 . 饮品定制弹窗 . 购物车 . 微信支付"),
       ("[攻略] 游玩攻略","5 条路线详情 . 难度/时长/景点/贴士/一键导航"),
       ("[美食] 美食探店","8 家周边美食深度攻略 . 必点菜 . 真实评价引用"),
       ("[服务] 快捷服务","21 项服务卡片式布局 . 点击即呼叫"),
       ("[地图] 地图导航","民宿精准定位 . 高德/腾讯地图跳转 . 交通指南"),
       ("[积分] 积分中心","积分余额 . 会员等级与升级进度条 . 兑换商城"),
       ("[订单] 订单管理","多平台订单统一列表 . 入住/退房/在住状态"),
       ("[后台] 员工看板","实时房态一览 . 服务请求追踪 . 经营数据概览")]

for nm,de in pages:
    p=DOC.add_paragraph(); pf=p.paragraph_format; pf.space_after=Pt(2); pf.space_before=Pt(6)
    r=p.add_run(nm); r.font.size=Pt(12); r.font.color.rgb=PRI; r.bold=True
    body(de,indent=True)

body("")
bullet("-", "所有页面响应式设计 . 适配各类手机屏幕尺寸")
bullet("-", "深色/浅色模式自动切换 . 跟随手机系统设置")
bullet("-", "春夏秋冬三季主题配色 . 页面视觉随季节自然变换")
bullet("-", "注：以上页面为项目规划内容，部分页面正在开发中")

DOC.add_page_break()

# ================================================================
# 尾页
# ================================================================
for _ in range(6):
    p=DOC.add_paragraph(); pf=p.paragraph_format; pf.space_after=Pt(0)

pe=DOC.add_paragraph(); pe.alignment=WD_ALIGN_PARAGRAPH.CENTER
re=pe.add_run("云上 . 归墅"); re.font.size=Pt(28); re.font.color.rgb=PRID; re.bold=True
ps=DOC.add_paragraph(); ps.alignment=WD_ALIGN_PARAGRAPH.CENTER
pfs=ps.paragraph_format; pfs.space_before=Pt(8)
rs=ps.add_run("云无心以出岫，鸟倦飞而知还"); rs.font.size=Pt(14); rs.font.color.rgb=ACC
pa=DOC.add_paragraph(); pa.alignment=WD_ALIGN_PARAGRAPH.CENTER
pfa=pa.paragraph_format; pfa.space_before=Pt(12)
ra=pa.add_run("庐山山上 . 大林沟路27号  |  16607927666"); ra.font.size=Pt(11); ra.font.color.rgb=GR
pv=DOC.add_paragraph(); pv.alignment=WD_ALIGN_PARAGRAPH.CENTER
rv=pv.add_run("微信公众号搜索「云上归墅民宿」体验全部功能\n项目规划方案  .  2026年6月"); rv.font.size=Pt(10); rv.font.color.rgb=LG

output=os.path.join(os.path.dirname(__file__),"云上归墅_项目规划方案_20260613.docx")
DOC.save(output)
print(f"云上归墅DOCX: {output}")
