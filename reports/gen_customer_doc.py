"""
面向民宿客户的功能方案介绍 — DOCX 格式
纯文本标记，不使用高位 Unicode emoji，确保所有系统正常渲染
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

DOC = Document()
for section in DOC.sections:
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)

PRIMARY = RGBColor(0x4A,0x7C,0x59); PRIMARY_DARK = RGBColor(0x3B,0x5E,0x47)
ACCENT = RGBColor(0xC8,0x96,0x6A); WHITE = RGBColor(0xFF,0xFF,0xFF)
DARK = RGBColor(0x2A,0x33,0x28); GRAY = RGBColor(0x5A,0x65,0x58)
LIGHT_GRAY = RGBColor(0x8A,0x95,0x86); BORDER = RGBColor(0xD8,0xE0,0xD4)

def section_title(icon_text, title):
    p1 = DOC.add_paragraph(); pf1 = p1.paragraph_format
    pf1.space_before = Pt(18); pf1.space_after = Pt(0)
    pPr = p1._p.get_or_add_pPr()
    pPr.append(parse_xml('<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="12" w:space="1" w:color="B8D0B8"/></w:pBdr>'))
    p2 = DOC.add_paragraph(); pf2 = p2.paragraph_format
    pf2.space_after = Pt(4); pf2.space_before = Pt(4)
    r = p2.add_run(f"[{icon_text}] {title}"); r.font.size = Pt(20)
    r.font.color.rgb = PRIMARY_DARK; r.bold = True
    p3 = DOC.add_paragraph(); pf3 = p3.paragraph_format; pf3.space_after = Pt(14)
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(parse_xml('<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="C8966A"/></w:pBdr>'))

def body(text, indent=False, size=11, color=DARK):
    p = DOC.add_paragraph(); pf = p.paragraph_format; pf.space_after = Pt(5)
    if indent: pf.left_indent = Cm(0.8)
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color

def bullet(icon, text, size=10.5):
    p = DOC.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(3); pf.left_indent = Cm(0.6)
    r = p.add_run(f"{icon} {text}"); r.font.size = Pt(size); r.font.color.rgb = DARK

def subsection(text, size=14):
    p = DOC.add_paragraph(); pf = p.paragraph_format
    pf.space_after = Pt(4); pf.space_before = Pt(10)
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = PRIMARY; r.bold = True

def _del_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        '<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    tblPr.append(parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders>'))

# ================================================================
#  封面
# ================================================================
for _ in range(4):
    p = DOC.add_paragraph(); pf = p.paragraph_format; pf.space_after = Pt(0)

p = DOC.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("民宿智能化运营方案"); r.font.size = Pt(30)
r.font.color.rgb = PRIMARY_DARK; r.bold = True
p2 = DOC.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf2 = p2.paragraph_format; pf2.space_after = Pt(6)
r2 = p2.add_run("微信公众号客服系统 . 网页全栈搭建"); r2.font.size = Pt(16); r2.font.color.rgb = ACCENT
p3 = DOC.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf3 = p3.paragraph_format; pf3.space_after = Pt(10)
r3 = p3.add_run("-" * 36); r3.font.size = Pt(10); r3.font.color.rgb = BORDER
p4 = DOC.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf4 = p4.paragraph_format; pf4.space_after = Pt(6)
r4 = p4.add_run("为精品民宿提供完整的线上智能服务解决方案\n覆盖微信公众号接入、AI管家、在线点单、游玩攻略、积分会员、多平台订单管理")
r4.font.size = Pt(11.5); r4.font.color.rgb = GRAY
DOC.add_paragraph()

p5 = DOC.add_paragraph(); p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf5 = p5.paragraph_format; pf5.space_before = Pt(12); pf5.space_after = Pt(4)
r5 = p5.add_run("参考案例：庐山 . 云上归墅民宿"); r5.font.size = Pt(13)
r5.font.color.rgb = PRIMARY_DARK; r5.bold = True

for c in ["11 间客房 | U 型三层山居小院 | 2016 年开业",
          "携程 4.7 分 . 85 条住客评价", "一楼配有三山二水咖啡馆", "全套系统已上线运行中"]:
    pc = DOC.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pfc = pc.paragraph_format; pfc.space_after = Pt(2)
    rc = pc.add_run(c); rc.font.size = Pt(10.5); rc.font.color.rgb = GRAY

DOC.add_paragraph()
p6 = DOC.add_paragraph(); p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run("2026年6月"); r6.font.size = Pt(10); r6.font.color.rgb = LIGHT_GRAY
DOC.add_page_break()

# ================================================================
#  1. 我们提供什么
# ================================================================
section_title("服务", "我们提供什么")
body("我们为民宿搭建一整套微信公众号智能客服系统 + 移动端 H5 网页。民宿只需拥有一个微信公众号，关注即用，住客无需下载任何 App。")
body("系统上线后，AI 智能管家全天候自动应答，在线点单、房型展示、攻略推荐、服务呼叫全部由系统自动完成，大幅降低前台人工回复量。")
body("")

subsection("一、核心能力")
for item in ["AI 智能管家全天候在线值守，无需轮班",
             "80% 以上的常规咨询由系统自动处理",
             "10 个功能页面覆盖民宿全场景服务",
             "携程 / 美团 / 飞猪 / 大众点评 4 大预订渠道打通"]:
    bullet("-", item)

DOC.add_paragraph()
subsection("二、交付内容")
deliverables = [
    ("微信公众号后端接入", "消息自动回复、关键词路由、AI 对话集成，关注即用"),
    ("移动端 H5 网页（10 个页面）", "房型展示、在线点单、游玩攻略、美食推荐、快捷服务、地图导航、积分中心、订单管理等"),
    ("AI 智能管家系统", "三阶段服务：预订前免费旅行顾问 - 预订后专属管家 - 离店后复购关怀"),
    ("积分会员体系", "积分获取/兑换/等级管理，增强客户粘性"),
    ("多平台订单管理看板", "携程/美团/飞猪/大众点评订单统一查看"),
    ("经营数据看板", "入住率、平台分布、收入趋势一目了然"),
]
for title, desc in deliverables:
    p = DOC.add_paragraph(); pf = p.paragraph_format; pf.space_after = Pt(2); pf.space_before = Pt(4)
    r = p.add_run(f">> {title}"); r.font.size = Pt(12); r.font.color.rgb = PRIMARY_DARK; r.bold = True
    body(desc, indent=True)

DOC.add_page_break()

# ================================================================
#  2. AI 智能管家
# ================================================================
section_title("AI", "AI 智能管家 -- 系统核心亮点")
body("这是整套系统中最能体现价值的部分。AI 管家不是简单的关键词回复机器人，而是能够根据住客所处的阶段，自动切换服务身份，提供恰如其分的帮助。")
body("对民宿的价值：一位 AI 管家可以同时服务所有住客，不会疲惫、不会遗漏、不会情绪化，为每一位客人提供一致的优质服务体验。")
body("")

phases = [
    ("阶段一：潜在客户 -- 免费旅行顾问",
     "任何关注公众号的用户都可以使用。AI 以本地旅行顾问的身份，热情解答景点、路线、交通、美食等问题。",
     ["零门槛使用，降低咨询心理负担", "自然地展示民宿特色和房型信息", "引导用户产生预订意愿", "即使最终不住宿，也留下了品牌好感"]),
    ("阶段二：已预订住客 -- 专属AI管家",
     "住客预订确认后自动解锁。AI 管家认识每位住客的姓名和偏好，提供个性化服务。",
     ["全天候即时响应，前台晚间也无需值守", "快捷服务一键呼叫：打扫/送餐/维修/叫醒等", "实时推荐周边美食和游玩路线", "大幅降低前台重复性咨询工作量"]),
    ("阶段三：已退房客人 -- 复购关怀",
     "离店后自动切换关怀模式，保持温暖的品牌连接。",
     ["退房后自动推送好评提醒，比人工催促更得体", "积分累计和升级通知，激励再次入住", "会员等级越高折扣越大，培养忠实回头客", "无需专人跟踪维护，系统自动运行"]),
]
for title, desc, items in phases:
    subsection(f">> {title}")
    body(desc, indent=True)
    for item in items: bullet("-", item)
    body("")

DOC.add_page_break()

# ================================================================
#  3. 页面功能详解
# ================================================================
section_title("功能", "移动端页面功能详解")
body("以下功能模块均以 H5 网页形式嵌入微信公众号菜单，住客点击即可访问。所有页面适配手机屏幕，支持浅色/深色模式自动切换。")
body("")

subsection("房型展示（3 个页面：列表 + 详情 + 图文轮播）")
body("8 种房型在线展示，每间房含多张实拍图片、面积、床型、设施列表和详细文字介绍。图片支持左右滑动轮播，价格信息透明展示。", indent=True)
bullet("-", "减少住客反复咨询「有没有XX房型」「多少钱」「能住几人」")
bullet("-", "引导住客到 OTA 平台直接预订，民宿无需自建支付系统")
bullet("-", "房型数据由民宿方随时更新维护")

subsection("在线点单（三山二水咖啡）")
body("住客在手机上浏览菜单、定制饮品偏好、下单支付。饮品支持糖度选择（全糖/半糖/无糖）、冰量选择（正常冰/少冰/温），云雾茶支持红茶/绿茶和杯/壶两档。下单后送餐到房。", indent=True)
bullet("-", "开拓非房费收入来源 -- 餐饮消费")
bullet("-", "无需额外开发 App，微信支付直接闭环")
bullet("-", "住客不用下楼即可点单，提升入住体验")

subsection("游玩攻略 + 美食推荐")
body("5 条深度游玩路线和 8 家周边美食探店推荐。每条路线和每家餐厅都有详细图文介绍，数据来源于小红书、大众点评、携程的真实住客探店笔记。", indent=True)
bullet("-", "大幅减少「附近有什么好玩的/好吃的」类重复咨询")
bullet("-", "内容真实有料，住客信任度高")
bullet("-", "美食攻略可直接跳转高德/腾讯地图导航")

subsection("快捷服务（21 项）")
body("住客在微信中回复关键词或点击页面按钮，即可呼叫客房服务、前台服务或设施维修。系统自动通知员工处理。", indent=True)
bullet("-", "客房：续住/打扫/补充用品/送餐/洗衣")
bullet("-", "前台：行李寄存/旅游咨询/叫车/叫醒/搭伙用餐/登山杖租借/特产代购等")
bullet("-", "维修：设施报修/空调/热水")
bullet("-", "服务请求自动通知员工，不遗漏")

subsection("积分会员体系")
body("住客入住消费、每日签到、写平台评价、邀请好友均可获取积分。积分可兑换咖啡、房型升级、延迟退房、房费抵扣券。会员分银卡/金卡/钻石卡三级，等级越高折扣越大。", indent=True)
bullet("-", "激励好评 -- 写评价 +80 分，直接提升 OTA 评分")
bullet("-", "激励回头 -- 积分累积 + 会员等级，复购动力")
bullet("-", "生日月 x1.5 倍积分 -- 温暖的小惊喜")

DOC.add_page_break()

# ================================================================
#  4. 运营看板
# ================================================================
section_title("管理", "员工后台 -- 运营数据一目了然")
body("为民宿运营方配备专属后台看板，无需登录多个 OTA 平台，一站式查看所有经营数据。")
body("")

subsection("多平台订单聚合看板")
bullet("-", "携程/美团/飞猪/大众点评/小红书/抖音/直接预订 -- 7 渠道订单统一管理")
bullet("-", "今日入住、今日退房、当前在住 -- 实时房态一目了然")
bullet("-", "未来 7 天到店预测 -- 合理安排人手和物资")
bullet("-", "各平台订单数和收入统计 -- 一目了然的渠道效果对比")

subsection("服务请求看板")
bullet("-", "住客通过公众号发起的每一条服务请求自动记录")
bullet("-", "待处理/处理中/已完成 -- 状态流转清晰")
bullet("-", "不遗漏任何一个客人需求")

subsection("自动化运营任务")
bullet("-", "退房后定时推送好评提醒 -- 无需人工跟踪")
bullet("-", "每周一自动生成经营周报 -- 数据驱动决策")
bullet("-", "多渠道口碑监控 -- 及时发现和回应评价")

DOC.add_page_break()

# ================================================================
#  5. 案例：云上归墅
# ================================================================
section_title("案例", "云上归墅民宿 -- 真实运行案例")
body("云上归墅是我们服务的首个民宿客户。以下为该系统在云上归墅的实际部署效果。")
body("")

subsection("民宿基本情况")
for item in ["位置：庐山山上 . 庐山风景名胜区大林沟路27号",
             "规模：11间客房，U型三层山居小院", "开业：2016年",
             "OTA评分：携程 4.7 分 . 85条评价",
             "配套：一楼三山二水咖啡馆", "预订渠道：携程/美团/飞猪/大众点评"]:
    bullet("-", item)

body("")
subsection("为云上归墅带来的价值")

values = [
    ("服务自动化", "住客的常规问题（房型、价格、路线、美食）由 AI 管家自动应答，前台只需处理少量个性化需求。即使民宿前台夜间休息，住客也能通过 AI 管家获得即时帮助。"),
    ("拓展收入来源", "三山二水咖啡馆线上点单系统，让住客在房间就能浏览菜单、定制饮品、下单支付。为咖啡简餐业务提供了便捷的销售渠道。"),
    ("提升好评率", "退房后系统自动推送得体、温暖的好评提醒，搭配积分激励（写评价 +80 分）。比人工催促的方式转化率更高，体验也更好。"),
    ("培养回头客", "积分体系 + 会员等级 + 生日月加成，让住客有持续回住的动力。从银卡到钻石卡，级别越高、专属福利越多。"),
    ("经营数据可视化", "老板通过手机就能查看实时入住情况、各平台订单分布、收入趋势。每周自动生成经营报告，辅助经营决策。"),
    ("品牌形象升级", "一个有 AI 管家、有在线点单、有深度攻略的公众号，让云上归墅在庐山民宿中脱颖而出。住客感知到的不是「一家民宿」，而是「一个品牌」。"),
]
for title, desc in values:
    subsection(f">> {title}")
    body(desc, indent=True)

DOC.add_page_break()

# ================================================================
#  6. 适用场景
# ================================================================
section_title("场景", "适用场景与民宿类型")
body("这套方案设计的初衷是为精品民宿、度假酒店、客栈青旅等中小型住宿业态提供智能化升级路径。")
body("")

subsection("特别适合以下场景")
scenarios = [
    ("精品民宿/设计师民宿", "房间数 5-30 间，希望提升品牌调性和住客体验，但无法负担自建 IT 团队"),
    ("景区周边民宿集群", "庐山、大理、莫干山、安吉等热门目的地，住客对周边游玩攻略需求强烈"),
    ("自带餐饮的民宿", "配有小咖啡馆、茶室、简餐厅的民宿，希望通过线上点单拓展非房收入"),
    ("多平台运营的民宿", "同时在携程/美团/飞猪/大众点评等渠道上线，需要一个统一的订单管理后台"),
    ("希望提升好评率的民宿", "有稳定的客流量但 OTA 评价数量偏少，需要系统化的评价激励和跟进机制"),
]
for title, desc in scenarios:
    subsection(f">> {title}")
    body(desc, indent=True)

DOC.add_page_break()

# ================================================================
#  7. 系统页面一览
# ================================================================
section_title("页面", "系统页面一览")

pages = [
    ("[首页] 首页", "民宿概览、精选房型推荐、AI 管家入口、季节主题切换"),
    ("[房型] 房型列表", "8 种房型卡片式展示，价格/人数/面积一目了然"),
    ("[详情] 房型详情", "多图轮播、设施清单、一键拨打预订电话"),
    ("[点单] 在线点单", "4 大品类菜单、饮品定制弹窗、购物车、微信支付"),
    ("[攻略] 游玩攻略", "5 条路线详情（难度/时长/景点/贴士/导航）"),
    ("[美食] 美食推荐", "8 家探店深度攻略、必点菜、真实评价、一键导航"),
    ("[服务] 快捷服务", "21 项服务卡片式展示，点击即呼叫"),
    ("[地图] 地图导航", "民宿定位、高德/腾讯地图跳转、交通指南"),
    ("[积分] 积分中心", "积分余额/会员等级/升级进度条/兑换商城/获取规则"),
    ("[订单] 订单管理", "多平台订单统一列表、状态筛选、详情查看"),
    ("[后台] 员工看板", "实时房态、服务请求、运营数据（内部页面）"),
]
for name, desc in pages:
    p = DOC.add_paragraph(); pf = p.paragraph_format; pf.space_after = Pt(2); pf.space_before = Pt(6)
    r = p.add_run(name); r.font.size = Pt(12); r.font.color.rgb = PRIMARY; r.bold = True
    body(desc, indent=True)

body("")
bullet("-", "深色/浅色模式自动适配，跟随手机系统设置")
bullet("-", "春夏秋冬三季主题配色切换，视觉随季节变化")
bullet("-", "所有页面均为响应式设计，适配不同手机屏幕尺寸")

DOC.add_page_break()

# ================================================================
#  8. 尾页
# ================================================================
for _ in range(6):
    p = DOC.add_paragraph(); pf = p.paragraph_format; pf.space_after = Pt(0)

p_end = DOC.add_paragraph(); p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run("让您的民宿也拥有智能大脑"); r_end.font.size = Pt(26)
r_end.font.color.rgb = PRIMARY_DARK; r_end.bold = True

p_sub = DOC.add_paragraph(); p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_sub = p_sub.paragraph_format; pf_sub.space_before = Pt(8)
r_sub = p_sub.add_run("微信公众号客服系统 . AI 管家 . 在线点单 . 会员积分 . 订单管理")
r_sub.font.size = Pt(12); r_sub.font.color.rgb = ACCENT

p_contact = DOC.add_paragraph(); p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf_contact = p_contact.paragraph_format; pf_contact.space_before = Pt(20)
r_contact = p_contact.add_run("参考案例：庐山 . 云上归墅民宿  |  微信公众号：云上归墅民宿")
r_contact.font.size = Pt(10.5); r_contact.font.color.rgb = GRAY

p_ver = DOC.add_paragraph(); p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ver = p_ver.add_run("方案版本 v2.14  .  2026年6月")
r_ver.font.size = Pt(9); r_ver.font.color.rgb = LIGHT_GRAY

# ================================================================
output = os.path.join(os.path.dirname(__file__), "云上归墅_民宿智能化运营方案_20260613.docx")
DOC.save(output)
print(f"DOCX 已保存至: {output}")
